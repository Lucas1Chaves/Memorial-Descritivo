import openpyxl
from openpyxl.utils import get_column_letter
import time
import docx
Not_Included = [14,15,16,17,18]
def Find_Description(Database,code,title):
    Result = {}
    # Primeiro tento achar a descrição com o código do item 
    try:
        index = Database['Codes'].index(code)
        Result['Description'] = Database['Descriptions'][index]
        Result['Method'] = 'Code'
        return Result
    except ValueError:
    # Caso não de certo, procuro a descrição com o titulo 
        try:
            index = Database['Titles'].index(title)
            Result['Description']= Database['Descriptions'][index]
            Result['Method'] = 'Title'
            return Result
        except ValueError:
            Result['Descriptions'] = ''
            Result['Method']='Missing'
            return Result
            

def String_To_Number(string):
    try:
        number = float(string)
    except:
        print("impossible to convert string to float, returning the string")
        return string
    if number.is_integer():
        number = int(number)
    return number

# Abrindo a tabela ORÇAMENTO dentro do arquivo Orçamento.xlsx
try:
    Budget_Sheet = openpyxl.load_workbook('Orçamento.xlsx')['ORÇAMENTO']
except:
    print('A arquivo do orçamento deve ter o nome "Orçamento.xlsx" e a tabela dentro do arquivo deve ter o nome "ORÇAMENTO"')
    time.sleep(15)
Budget = {}
Database = {}
# Obtendo a primeira coluna de celulas e o itens dela
Column_A = Budget_Sheet['A2':'A'+str(Budget_Sheet.max_row)]
Budget['Items'] = [String_To_Number(cell[0].value) for cell in Column_A]

# Obtendo a segunda coluna de celulas e os codigos
Column_B = Budget_Sheet['B2':'B'+str(Budget_Sheet.max_row)]
Budget['Codes'] = [cell[0].value for cell in Column_B]

# Obtendo a terceira coluna de celulas e os titulos
Column_C = Budget_Sheet['C2':'C'+str(Budget_Sheet.max_row)]
Budget['Titles'] = [cell[0].value for cell in Column_C]

# Abrindo a primeira tabela do arquivo de Base de Dados
Database_Sheet = openpyxl.load_workbook('Base de Dados.xlsx')['Plan1']
# Obtendo a primeira coluna de celulas e os codigos
Column_A = Database_Sheet['A2':'A'+str(Database_Sheet.max_row)]
Database['Codes'] = [cell[0].value for cell in Column_A]

# Obtendo a segunda coluna de celulas e os titulos
Column_B = Database_Sheet['B2':'B'+str(Database_Sheet.max_row)]
Database['Titles'] = [cell[0].value for cell in Column_B]

# Obtendo a terceira coluna de celulas e as descrições
Column_C = Database_Sheet['C2':'C'+str(Database_Sheet.max_row)]
Database['Descriptions'] = [cell[0].value for cell in Column_C]


doc = docx.Document('Papel Timbrado Sião.docx')
doc_missing_codes = docx.Document()
#font = doc.styles['Normal'].font
#font.name = 'Arial'
missing_codes =[]
for item,title,code in zip(Budget['Items'],Budget['Titles'],Budget['Codes']):
    
    if item != '' and int(item) not in (Not_Included):
        print(code,item,type(item))
        # Se o código não existe e o item é um inteiro,  
        # então significa que é um titulo 
        if code ==None and isinstance(item,int):
            
            doc.add_paragraph(str(item)+'. ' +title,style='01-TÍTULO')
            #p.paragraph_format.space_after = docx.shared.Pt(0)
            #r=p.add_run(str(item)+'. ' +title)
        
            #r.font.highlight_color = docx.enum.text.WD_COLOR.GRAY_25
            
            #r.bold =True
        
        else:
            doc.add_paragraph(str(item) +' '+ title,style='02-SERVIÇO')
            #p2.paragraph_format.space_after = docx.shared.Pt(0)
            #p2.add_run(item.strip() +' '+ titulo).bold = True
            Result = Find_Description(Database,code,title)
            if Result['Method'] == 'Code':
                doc.add_paragraph(Result['Description'],style='03-DESCRIÇÃO')
            elif Result['Method'] == 'Title':
                doc.add_paragraph(Result['Description'],style='03-DESCRIÇÃO-TITULO')

            else:
                doc_missing_codes.add_paragraph(str(code))
        
            

doc.save('Memorial Descritivo1.docx')
doc_missing_codes.save('Itens Sem Código.docx')
print('Deu Tudo certo :) ')
time.sleep(10)